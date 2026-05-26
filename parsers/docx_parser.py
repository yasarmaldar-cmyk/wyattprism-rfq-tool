import docx
import io
from .document import ParsedDocument


def _docx_table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""

    headers = rows[0]
    md = "| " + " | ".join(headers) + " |\n"
    md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    for row in rows[1:]:
        while len(row) < len(headers):
            row.append("")
        row = row[: len(headers)]
        md += "| " + " | ".join(row) + " |\n"
    return md


def parse_docx(file_bytes: bytes, filename: str = "document.docx") -> ParsedDocument:
    doc = docx.Document(io.BytesIO(file_bytes))
    tables = []
    text_parts = []

    # Track which tables we've seen by their position relative to paragraphs
    table_index = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = element
            # Get the paragraph text
            text = "".join(node.text or "" for node in para.iter() if node.text)
            # Check for heading style
            style = para.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
            )
            if style is not None:
                style_val = style.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                    "",
                )
                if "Heading" in style_val or "heading" in style_val:
                    level = "".join(c for c in style_val if c.isdigit()) or "1"
                    text = f"\n{'#' * int(level)} {text}"
            if text.strip():
                text_parts.append(text.strip())

        elif tag == "tbl" and table_index < len(doc.tables):
            table = doc.tables[table_index]
            md = _docx_table_to_markdown(table)
            if md:
                tables.append({"page": 0, "markdown": md})
                text_parts.append("\n[Table]\n" + md)
            table_index += 1

    raw_text = "\n\n".join(text_parts)

    return ParsedDocument(
        raw_text=raw_text,
        pages=[raw_text],  # DOCX doesn't have pages
        tables=tables,
        filename=filename,
        page_count=1,
        file_type="docx",
    )
