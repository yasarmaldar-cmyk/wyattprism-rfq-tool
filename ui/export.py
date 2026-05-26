import io
import json
import docx
from pathlib import Path
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import streamlit as st

LOGO_PATH = Path(__file__).parent.parent / "assets" / "logo.png"
GRADIENT_PATH = Path(__file__).parent.parent / "assets" / "gradient_bar.png"
GRADIENT_REV_PATH = Path(__file__).parent.parent / "assets" / "gradient_bar_rev.png"

# WP Brand colors
WP_CORAL = RGBColor(232, 77, 61)       # #E84D3D
WP_RED = RGBColor(217, 68, 82)         # #D94452
WP_MAGENTA = RGBColor(139, 47, 143)    # #8B2F8F
WP_DARK = RGBColor(42, 21, 64)         # #2A1540
WP_LIGHT_BG = RGBColor(250, 245, 252)  # #FAF5FC
WP_GRAY = RGBColor(100, 100, 100)
WP_WHITE = RGBColor(255, 255, 255)


def export_clarifications_docx(clarifications: list, filename: str) -> bytes:
    """Export just the clarification questions as a standalone DOCX."""
    doc = Document()

    doc.add_heading("Clarification Questions", level=0)
    doc.add_paragraph(f"Document: {filename}")
    doc.add_paragraph("")

    if not isinstance(clarifications, list) or not clarifications:
        doc.add_paragraph("No clarification questions generated.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # Framework questions first
    framework_qs = [q for q in clarifications if q.get("category", "").lower() == "framework"]
    other_qs = [q for q in clarifications if q.get("category", "").lower() != "framework"]

    if framework_qs:
        doc.add_heading("Framework & Reporting Standards", level=1)
        for i, q in enumerate(framework_qs, 1):
            severity = q.get("severity", "").upper()
            p = doc.add_paragraph()
            run = p.add_run(f"Q{i}. [{severity}] ")
            run.bold = True
            if severity == "HIGH":
                run.font.color.rgb = RGBColor(220, 50, 50)
            elif severity == "MEDIUM":
                run.font.color.rgb = RGBColor(230, 150, 0)
            else:
                run.font.color.rgb = RGBColor(50, 160, 50)
            p.add_run(q.get("question", ""))
            if q.get("reason"):
                reason_p = doc.add_paragraph()
                reason_p.paragraph_format.left_indent = Pt(20)
                run = reason_p.add_run(f"Why: {q.get('reason', '')}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)

    num_offset = len(framework_qs)
    if other_qs:
        doc.add_heading("Other Clarifications", level=1)
        for i, q in enumerate(other_qs, num_offset + 1):
            severity = q.get("severity", "").upper()
            category = q.get("category", "general").upper()
            p = doc.add_paragraph()
            run = p.add_run(f"Q{i}. [{severity} | {category}] ")
            run.bold = True
            if severity == "HIGH":
                run.font.color.rgb = RGBColor(220, 50, 50)
            elif severity == "MEDIUM":
                run.font.color.rgb = RGBColor(230, 150, 0)
            else:
                run.font.color.rgb = RGBColor(50, 160, 50)
            p.add_run(q.get("question", ""))
            if q.get("reason"):
                reason_p = doc.add_paragraph()
                reason_p.paragraph_format.left_indent = Pt(20)
                run = reason_p.add_run(f"Why: {q.get('reason', '')}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_colored_border_bottom(paragraph, color_hex="E84D3D"):
    """Add a colored bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="4" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_shading(paragraph, color_hex="FAF5FC"):
    """Add background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    pPr.append(shd)


def _add_gradient_bar(doc, reverse=False):
    """Add a smooth gradient bar as an image — real gradient, not table columns."""
    img_path = GRADIENT_REV_PATH if reverse else GRADIENT_PATH
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str(img_path), width=Cm(17))


def _add_left_border(paragraph, color_hex="8B2F8F"):
    """Add a colored left border accent to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def export_proposal_docx(proposal_text: str, client_name: str, report_type: str) -> bytes:
    """Export the proposal draft as a professionally branded DOCX."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(51, 51, 51)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === HEADER WITH LOGO ===
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(4))

    # === SMOOTH GRADIENT BAR ===
    _add_gradient_bar(doc)

    # === PARSE AND RENDER PROPOSAL TEXT ===
    lines = proposal_text.strip().split("\n")
    section_keywords = {
        "Scope of Work": "D94452",
        "Scope of Digital": "D94452",
        "Please note": "8B2F8F",
        "Please Note": "8B2F8F",
        "Timeline": "E84D3D",
        "Costing": "D94452",
        "Cost": "D94452",
        "Payment terms": "8B2F8F",
        "Payment Terms": "8B2F8F",
        "Benefits": "E84D3D",
        "Additional Scope": "8B2F8F",
        "Digital E-Report Exclusions": "8B2F8F",
        "Digital E-Report Timeline": "E84D3D",
        "Digital E-Report Cost": "D94452",
        "How It Works": "D94452",
        "Scope includes": "D94452",
    }

    # Sub-headers for optional add-ons (styled differently — smaller, with left accent)
    addon_headers = [
        "Eco-Friendly Digital Report",
        "Animated Flipbook",
        "RAG AI",
        "120-Second Highlights Video",
    ]

    prev_was_blank = False
    for line in lines:
        stripped = line.strip()

        # Skip consecutive blank lines — use paragraph spacing instead
        if not stripped:
            if not prev_was_blank:
                prev_was_blank = True
            continue
        prev_was_blank = False
        stripped_lower = stripped.lower()

        # --- Appendix separator ---
        if stripped.startswith("---") and "APPENDIX" in stripped.upper():
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)
            if LOGO_PATH.exists():
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p2.paragraph_format.space_after = Pt(8)
                run2 = p2.add_run()
                run2.add_picture(str(LOGO_PATH), width=Cm(3))
            _add_gradient_bar(doc)
            continue

        # --- Appendix title: "Digital E-Report for..." ---
        if stripped.startswith("Digital E-Report for") or stripped.startswith("Digital Report for"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = WP_MAGENTA
            continue

        # --- "As a value-add" intro line ---
        if stripped.startswith("As a value-add"):
            p = doc.add_paragraph()
            _add_shading(p, "F8F0FA")
            _add_left_border(p, "8B2F8F")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped)
            run.font.size = Pt(11)
            run.italic = True
            run.font.color.rgb = WP_MAGENTA
            continue

        # --- Title: "Proposal for..." ---
        if stripped.startswith("Proposal for"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = WP_DARK
            continue

        # --- Date line ---
        if "|" in stripped and any(m in stripped.lower() for m in ["january", "february", "march", "april", "may", "june", "july", "august", "september", "october", "november", "december"]):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            run.font.color.rgb = WP_GRAY
            continue

        # --- Dear Sir/Madam ---
        if stripped_lower.startswith("dear "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(stripped)
            run.font.size = Pt(11)
            run.italic = True
            continue

        # --- "Part A:", "Part B:" sub-section headers ---
        if stripped_lower.startswith("part ") and ":" in stripped[:20]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped.rstrip(":"))
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = WP_DARK
            _add_left_border(p, "E84D3D")
            continue

        # --- "Client to provide:" info line ---
        if stripped_lower.startswith("client to provide"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            _add_shading(p, "FFF8F5")
            _add_left_border(p, "E84D3D")
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            run.font.color.rgb = WP_GRAY
            run.italic = True
            continue

        # --- Optional add-on sub-headers (check before section headers) ---
        is_addon = False
        for ah in addon_headers:
            if stripped_lower.startswith(ah.lower()) and stripped.rstrip().endswith(":"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                _add_left_border(p, "D94452")
                _add_shading(p, "FFF8F5")
                run = p.add_run(stripped.rstrip(":"))
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = WP_RED
                is_addon = True
                break
        if is_addon:
            continue

        # --- Section headers (case-insensitive match) ---
        is_section = False
        for kw, color in section_keywords.items():
            if stripped_lower.startswith(kw.lower()):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                # Use title case for display, strip colon
                header_text = stripped.rstrip(":").strip()
                run = p.add_run(header_text)
                run.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])
                _add_colored_border_bottom(p, color)
                is_section = True
                break
        if is_section:
            continue

        # --- Bullet points ---
        if stripped.startswith(("-", "\u2022", "\u2013")):
            clean = stripped.lstrip("-\u2022\u2013 ")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("\u25B8 ")
            run.font.color.rgb = WP_CORAL
            run.font.size = Pt(10)
            run2 = p.add_run(clean)
            run2.font.size = Pt(10.5)
            continue

        # --- Costing line ---
        if ("Rs." in stripped or "\u20b9" in stripped) and ("charges" in stripped.lower() or "cost" in stripped.lower() or "lac" in stripped.lower() or "discussed" in stripped.lower()):
            p = doc.add_paragraph()
            _add_shading(p, "FFF5F5")
            _add_left_border(p, "D94452")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = WP_RED
            continue

        # --- Digital teaser callout box ---
        if stripped_lower.startswith(">>>") or (stripped_lower.startswith("digital proposition") and "see next page" in stripped_lower):
            clean_text = stripped.lstrip("> ").strip()
            # Split "DIGITAL PROPOSITION:" from the rest
            if ":" in clean_text:
                label, body = clean_text.split(":", 1)
            else:
                label, body = "DIGITAL PROPOSITION", clean_text

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(14)
            _add_shading(p, "2A1540")
            _add_left_border(p, "E84D3D")

            # Label in coral
            run1 = p.add_run(f"  {label.strip()}:  ")
            run1.bold = True
            run1.font.size = Pt(11)
            run1.font.color.rgb = WP_CORAL

            # Body in white
            run2 = p.add_run(body.strip())
            run2.font.size = Pt(10.5)
            run2.font.color.rgb = WP_WHITE

            # Arrow indicator
            run3 = p.add_run("  \u2192")
            run3.bold = True
            run3.font.size = Pt(12)
            run3.font.color.rgb = WP_CORAL
            continue

        # --- WP assures closing ---
        if stripped_lower.startswith("wp assures") or stripped_lower.startswith("wyattprism assures"):
            p = doc.add_paragraph()
            _add_shading(p, "F8F5FA")
            _add_left_border(p, "8B2F8F")
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped)
            run.font.size = Pt(11)
            continue

        # --- Sincerely ---
        if stripped_lower.startswith("sincerely"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            run = p.add_run(stripped)
            run.italic = True
            run.font.size = Pt(11)
            continue

        # --- Signatory ---
        if stripped.startswith("CA ") or stripped.startswith("Shailesh"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = WP_DARK
            continue

        if stripped_lower.startswith("director") or stripped_lower.startswith("president"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            run.font.color.rgb = WP_GRAY
            continue

        if stripped.startswith("+91"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            run.font.color.rgb = WP_CORAL
            continue

        # --- "How It Works" / "Scope includes" ---
        if stripped_lower.startswith("how it works") or stripped_lower.startswith("scope includes"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped.rstrip(":"))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = WP_DARK
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(stripped)
        run.font.size = Pt(11)

    # === FOOTER SMOOTH GRADIENT (reversed) ===
    _add_gradient_bar(doc, reverse=True)

    # Confidential footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("CONFIDENTIAL | WyattPrism Communications Pvt Ltd")
    run.font.size = Pt(8)
    run.font.color.rgb = WP_GRAY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _safe(text, max_len=500):
    if not text:
        return "N/A"
    s = str(text)[:max_len]
    # Replace characters that cause encoding issues in PDF
    s = s.encode("latin-1", errors="replace").decode("latin-1")
    # Remove leading whitespace that eats into cell width
    return s.strip()


def _safe_multi_cell(pdf, w, h, text):
    """Safe multi_cell that won't crash on edge cases."""
    try:
        pdf.multi_cell(w, h, text)
    except Exception:
        # Fallback: truncate and retry
        try:
            pdf.multi_cell(w, h, text[:100] + "...")
        except Exception:
            pdf.cell(w, h, "...", ln=True)


def export_pdf(results: dict, filename: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Cover page
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(0, 40, "RFQ Analysis Report", ln=True, align="C")
    pdf.set_font("Helvetica", "", 14)
    pdf.cell(0, 10, _safe(filename), ln=True, align="C")
    pdf.ln(20)

    summary = results.get("summary", {})
    org = summary.get("issuing_organization", {})
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, f"Organization: {_safe(org.get('name', 'N/A'))}", ln=True)
    pdf.cell(0, 8, f"Document Type: {_safe(summary.get('document_type', 'N/A'))}", ln=True)
    pdf.cell(0, 8, f"Report Type: {_safe(summary.get('report_type', 'N/A'))}", ln=True)

    # Frameworks
    frameworks = summary.get("frameworks", {})
    detected = frameworks.get("detected", [])
    if detected:
        pdf.cell(0, 8, f"Frameworks: {_safe(', '.join(detected))}", ln=True)
    elif frameworks.get("no_framework_detected"):
        pdf.set_text_color(200, 0, 0)
        pdf.cell(0, 8, "WARNING: No reporting framework specified", ln=True)
        pdf.set_text_color(0, 0, 0)

    # Summary section
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Summary", ln=True)
    pdf.set_font("Helvetica", "", 10)

    scope = summary.get("scope_of_work", {})
    _safe_multi_cell(pdf, 0, 6, f"Scope: {_safe(scope.get('description', 'N/A'), 500)}")
    pdf.ln(5)

    if scope.get("deliverables"):
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, "Deliverables:", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for d in scope["deliverables"]:
            pdf.cell(0, 6, f"  - {_safe(d)}", ln=True)

    # Timeline
    timeline = summary.get("timeline", {})
    if timeline.get("key_dates"):
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, "Key Dates:", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for d in timeline["key_dates"]:
            pdf.cell(0, 6, f"  - {_safe(d.get('event', ''))}: {_safe(d.get('date', ''))}", ln=True)

    # Clarifications
    clarifications = results.get("clarifications", [])
    if isinstance(clarifications, list) and clarifications:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "Clarification Questions", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for i, q in enumerate(clarifications, 1):
            severity = q.get("severity", "").upper()
            pdf.set_font("Helvetica", "B", 10)
            _safe_multi_cell(pdf, 0, 6, f"{i}. [{severity}] {_safe(q.get('question', ''), 200)}")
            pdf.set_font("Helvetica", "", 9)
            reason = _safe(q.get("reason", ""), 200)
            if reason and reason != "N/A":
                _safe_multi_cell(pdf, 0, 5, f"Reason: {reason}")
            pdf.ln(3)

    # Recommended frameworks (if no framework detected)
    recommended = frameworks.get("recommended_frameworks", [])
    if recommended:
        pdf.ln(5)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 8, "Recommended Frameworks:", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for fw in recommended:
            pdf.cell(0, 6, f"  - {_safe(fw)}", ln=True)

    return bytes(pdf.output())


def export_docx(results: dict, filename: str) -> bytes:
    doc = Document()

    # Title
    title = doc.add_heading("RFQ Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(filename, style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = results.get("summary", {})
    org = summary.get("issuing_organization", {})

    doc.add_paragraph(f"Organization: {org.get('name', 'N/A')}")
    doc.add_paragraph(f"Document Type: {summary.get('document_type', 'N/A')}")
    doc.add_paragraph(f"Report Type: {summary.get('report_type', 'N/A')}")

    # Frameworks
    frameworks = summary.get("frameworks", {})
    detected = frameworks.get("detected", [])
    if detected:
        doc.add_paragraph(f"Frameworks: {', '.join(detected)}")
    elif frameworks.get("no_framework_detected"):
        p = doc.add_paragraph()
        run = p.add_run("WARNING: No reporting framework specified")
        run.font.color.rgb = RGBColor(200, 0, 0)
        run.bold = True

    # Summary
    doc.add_heading("Summary", level=1)
    scope = summary.get("scope_of_work", {})
    doc.add_paragraph(scope.get("description", "N/A"))

    if scope.get("deliverables"):
        doc.add_heading("Deliverables", level=2)
        for d in scope["deliverables"]:
            doc.add_paragraph(d, style="List Bullet")

    # Timeline
    timeline = summary.get("timeline", {})
    if timeline.get("key_dates"):
        doc.add_heading("Key Dates", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Event"
        hdr[1].text = "Date"
        hdr[2].text = "Notes"
        for d in timeline["key_dates"]:
            row = table.add_row().cells
            row[0].text = str(d.get("event", ""))
            row[1].text = str(d.get("date", ""))
            row[2].text = str(d.get("notes", "") or "")

    # Clarifications
    clarifications = results.get("clarifications", [])
    if isinstance(clarifications, list) and clarifications:
        doc.add_heading("Clarification Questions", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "Severity"
        hdr[2].text = "Question"
        hdr[3].text = "Reason"
        for i, q in enumerate(clarifications, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = q.get("severity", "").upper()
            row[2].text = q.get("question", "")
            row[3].text = q.get("reason", "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_export():
    results = st.session_state.get("analysis_results")
    if not results:
        return

    filename = st.session_state.get("uploaded_filename", "document")

    st.sidebar.divider()
    st.sidebar.subheader("Export")

    col1, col2 = st.sidebar.columns(2)
    with col1:
        pdf_bytes = export_pdf(results, filename)
        st.download_button(
            "Download PDF",
            data=pdf_bytes,
            file_name=f"{filename}_analysis.pdf",
            mime="application/pdf",
        )
    with col2:
        docx_bytes = export_docx(results, filename)
        st.download_button(
            "Download DOCX",
            data=docx_bytes,
            file_name=f"{filename}_analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
